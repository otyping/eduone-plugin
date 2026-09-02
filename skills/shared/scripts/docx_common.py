# -*- coding: utf-8 -*-
"""
docx_common.py — ไลบรารีกลางสำหรับสร้าง .docx ของโปรเจกต์ EDU ONE (สพฐ.)

รวม helper ที่ build_content.py / build_lesson_plan.py / build_exercise.py ใช้ร่วมกัน:
- ฟอนต์ TH Sarabun New (latin + complex-script ไทย) + ขนาด szCs
- inline markup: **bold**, ^{superscript}, _{subscript}, \\alias (จาก symbols.txt)
- LaTeX -> OMML (Office Math) สำหรับสมการจริง (ถ้ามี toolchain)
- ตาราง 2 คอลัมน์ (หน้าปก/แผน) + ตารางเนื้อหาอิสระ
- หน้ากระดาษ A4, header ขวา, footer เลขหน้า
- Thai line-break: แทรก ZWSP ตามขอบคำ (pythainlp newmm) + ตั้ง w:lang bidi=th-TH

พอร์ตจากโปรเจกต์ OVEC (build_docx.py ของ lesson-plan + exercise) แล้วรวมเป็นไลบรารีเดียว
รันผ่านตัวห่อ: eduone-py <ชื่อไฟล์นี้> <args>  (หา Python 3.12 ให้เองทุก OS)
"""
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import math_policy as mp  # noqa: E402

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Mm, RGBColor

try:
    from pythainlp.tokenize import word_tokenize as _th_tokenize
except Exception:  # ไม่มี pythainlp ก็ข้าม ZWSP (ยังตั้ง lang ให้)
    _th_tokenize = None

FONT_NAME = "TH Sarabun New"
BODY_SIZE = Pt(14)
HEADING_SIZE = Pt(16)
COVER_TITLE_SIZE = Pt(18)
HF_SIZE = Pt(12)
BLACK = RGBColor(0, 0, 0)
# สีสำหรับแบบฝึกหัดที่มีสื่อ (ฉบับผู้ตรวจ):
#   BLUE = บทเสียง (ให้คนตรวจรู้ว่าเสียงที่ได้ยินพูดว่าอะไร)
#   RED  = ตัวเลือกที่เป็นคำตอบถูก
BLUE = RGBColor(0x1F, 0x4E, 0x9C)
RED = RGBColor(0xC0, 0x00, 0x00)
ZWSP = "​"
_THAI_RE = re.compile(r"[฀-๿]")

# --- inline-syntax: LaTeX aliases (\\Omega ฯลฯ) ---
# ย้ายไปอยู่ math_policy.py แล้ว เพราะ .pptx ต้องใช้ชุดเดียวกัน — คงชื่อเดิมไว้ให้
# โค้ดที่เรียก dc._apply_aliases / dc.SYMBOLS_PATH ใช้ได้เหมือนเดิม
SYMBOLS_PATH = mp.SYMBOLS_PATH
_ALIAS_RE = mp._ALIAS_RE
_ALIASES = mp._ALIASES
_apply_aliases = mp.apply_aliases


def _parse_inline_stateful(text, bold=False):
    """คืน (runs, bold ที่ค้างอยู่ท้ายข้อความ) — runs = (segment, bold, vert)

    ★ ต้องรับ/คืนสถานะ bold เพราะข้อความหนึ่งย่อหน้าถูกตัดเป็นหลายท่อนก่อนถึงที่นี่
      (ท่อนสมการ `$...$` คั่นอยู่) ถ้าไม่ส่งสถานะต่อ คู่ `**` ที่คร่อมสมการจะเพี้ยน
    """
    text = _apply_aliases(text)
    runs = []
    buf = []
    i = 0

    def flush(vert=None):
        if buf:
            runs.append(("".join(buf), bold, vert))
            buf.clear()

    while i < len(text):
        if text[i:i + 2] == "**":
            flush()
            bold = not bold
            i += 2
            continue
        if text[i:i + 2] == "^{":
            close = text.find("}", i + 2)
            if close != -1:
                flush()
                inner = text[i + 2:close]
                if inner:
                    runs.append((inner, bold, "superscript"))
                i = close + 1
                continue
        if text[i:i + 2] == "_{":
            close = text.find("}", i + 2)
            if close != -1:
                flush()
                inner = text[i + 2:close]
                if inner:
                    runs.append((inner, bold, "subscript"))
                i = close + 1
                continue
        buf.append(text[i])
        i += 1
    flush()
    return runs, bold


def _parse_inline(text):
    """คืน list ของ (segment, bold, vert) — vert in {None,'superscript','subscript'}"""
    return _parse_inline_stateful(text)[0]


# --- LaTeX -> OMML (สมการ Word จริง) ---
try:
    import glob as _glob
    from lxml import etree as _etree
    from latex2mathml.converter import convert as _latex2mathml

    _XSL = os.path.join(os.path.dirname(__file__), "MML2OMML.XSL")
    if not os.path.exists(_XSL):
        _c = _glob.glob(r"C:\Program Files*\Microsoft Office\root\Office*\MML2OMML.XSL")
        _XSL = _c[0] if _c else None
    _OMML_XSLT = _etree.XSLT(_etree.parse(_XSL)) if _XSL else None
except Exception:
    _OMML_XSLT = None
    _latex2mathml = None

_UNI2TEX = {"×": r"\times ", "÷": r"\div ", "−": "-", "·": r"\cdot ",
            "≈": r"\approx ", "≤": r"\leq ", "≥": r"\geq ", "≠": r"\neq ",
            "±": r"\pm ", "°": r"\degree ", "Ω": r"\Omega ", "π": r"\pi ",
            "µ": r"\mu ", "μ": r"\mu "}

_MATH_MARKER_RE = re.compile(r"\\frac|\\sqrt|\^\{|_\{|\\[A-Za-z]+")
_NONTHAI_RE = re.compile(r"[^฀-๿]+")
_NUMLABEL_RE = re.compile(r"^\d+\.\s+")
_PUNCT_LEAD_RE = re.compile(r"^[\s.,:;]*")
_PUNCT_TRAIL_RE = re.compile(r"[\s.,:;]*$")


def latex_to_omml(latex):
    if _OMML_XSLT is None or _latex2mathml is None:
        return None
    try:
        s = latex
        for u, t in _UNI2TEX.items():
            if u in s:
                s = s.replace(u, t)
        mathml = _latex2mathml(s)
        return _OMML_XSLT(_etree.fromstring(mathml.encode("utf-8"))).getroot()
    except Exception:
        return None


def _auto_spans(text):
    out = []
    pos = 0
    for m in _NONTHAI_RE.finditer(text):
        if m.start() > pos:
            out.append(("text", text[pos:m.start()]))
        chunk = m.group(0)
        if _MATH_MARKER_RE.search(chunk):
            lead = ""
            ml = _NUMLABEL_RE.match(chunk)
            if ml:
                lead = chunk[:ml.end()]; chunk = chunk[ml.end():]
            lp = _PUNCT_LEAD_RE.match(chunk).end()
            lead += chunk[:lp]; chunk = chunk[lp:]
            tp = _PUNCT_TRAIL_RE.search(chunk).start()
            trail, core = chunk[tp:], chunk[:tp]
            if lead:
                out.append(("text", lead))
            if core:
                out.append(("math", core))
            if trail:
                out.append(("text", trail))
        else:
            out.append(("text", chunk))
        pos = m.end()
    if pos < len(text):
        out.append(("text", text[pos:]))
    return out


def _split_math_spans(text):
    """$...$ = บังคับสมการ; นอกนั้น auto-detect"""
    out = []
    pos = 0
    for m in re.finditer(r"\$(.+?)\$", text):
        if m.start() > pos:
            out.extend(_auto_spans(text[pos:m.start()]))
        out.append(("math", m.group(1)))
        pos = m.end()
    if pos < len(text):
        out.extend(_auto_spans(text[pos:]))
    return out


def render_spans(text, math=True):
    """คืน list ของ (kind, value, bold, vert) — kind in {'text','math'}

    ★ ลำดับสำคัญมาก และเคยผิดมาก่อน
      1. ตัด `$...$` ออกก่อน — ในนั้นเป็น LaTeX ห้ามใครแตะ (`x^{2}` ต้องอยู่ครบ)
      2. ท่อนที่เหลือค่อยอ่าน markup `**ตัวหนา**` `^{}` `_{}` โดยส่งสถานะ bold ต่อกัน
      3. ท่อนข้อความธรรมดาค่อย auto-detect สมการที่ไม่ได้ใส่ `$...$`

      เดิมทำข้อ 3 ก่อนข้อ 2 — `_NONTHAI_RE` นับ `**` เป็นอักขระไม่ใช่ไทย จึงตัดคู่
      `**...**` ขาดจากกัน แล้วเรียก `_parse_inline` แยกทีละท่อนโดยรีเซ็ตสถานะทุกครั้ง
      ผลคือ `**จำนวนนับ** (counting number)` ได้ตัวหนาไปลงที่ `(counting number)`
      ส่วนคำไทยที่ควรหนากลับไม่หนา (พบในเอกสารที่ผลิตจริงทุกไฟล์)
    """
    if not math:
        return [("text", seg, b, v) for seg, b, v in _parse_inline(text)]

    out = []

    def emit_text(chunk, bold):
        if not chunk:
            return bold
        runs, bold = _parse_inline_stateful(chunk, bold)
        for seg, seg_bold, vert in runs:
            if vert:                     # ตัวยก/ตัวห้อยเป็น run ข้อความ ไม่ใช่สมการ
                out.append(("text", seg, seg_bold, vert))
                continue
            for kind, val in _auto_spans(seg):
                out.append((kind, val, seg_bold, None))
        return bold

    bold = False
    pos = 0
    for m in re.finditer(r"\$(.+?)\$", text):
        bold = emit_text(text[pos:m.start()], bold)
        out.append(("math", m.group(1), bold, None))
        pos = m.end()
    emit_text(text[pos:], bold)
    return out


def put_math(paragraph, latex, size, bold=False, color=None, italic=False):
    """เขียนสมการหนึ่งช่วงลงย่อหน้า ตามนโยบายใน math_policy

    ซ้อนชั้น (เศษส่วน/กรณฑ์/เมทริกซ์) -> วัตถุ Equation จริง
    ไม่ซ้อนชั้น (ตัวเลข เครื่องหมาย ยกกำลังชั้นเดียว) -> run ข้อความธรรมดา
      ซึ่งคัดลอกได้ ค้นหาเจอ แก้ในเวิร์ดได้ และไม่ดันความสูงบรรทัด

    ถ้าเครื่องไม่มี toolchain แปลง OMML จะตกมาทางข้อความให้เอง ไม่ทิ้งสูตรหาย
    """
    if mp.is_hard(latex):
        om = latex_to_omml(latex)
        if om is not None:
            paragraph._p.append(om)
            return True
    for part, vert, ital in mp.simple_runs(latex):
        run = paragraph.add_run(part)
        set_run_font(run, size, bold=bold, color=color, italic=italic or ital)
        _set_vert_align(run, vert)
    return False


def _set_vert_align(run, vert):
    if not vert:
        return
    rpr = run._element.get_or_add_rPr()
    va = rpr.find(qn("w:vertAlign"))
    if va is None:
        va = OxmlElement("w:vertAlign")
        rpr.append(va)
    va.set(qn("w:val"), vert)


def set_run_font(run, size, bold=False, color=None, italic=False):
    """ตั้งฟอนต์ TH Sarabun New ทั้ง latin และ complex-script (ไทย)

    color=None -> สีดำ (พฤติกรรมเดิม); ส่ง RGBColor เพื่อใช้สีอื่น (BLUE/RED)
    """
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    if italic:
        run.font.italic = True
    run.font.color.rgb = BLACK if color is None else color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT_NAME)
    szcs = rpr.find(qn("w:szCs"))
    if szcs is None:
        szcs = OxmlElement("w:szCs")
        rpr.append(szcs)
    szcs.set(qn("w:val"), str(int(size.pt * 2)))


def clear_paragraph(paragraph):
    for r in list(paragraph.runs):
        r._r.getparent().remove(r._r)
    return paragraph


def set_para_align(paragraph, jc="thaiDistribute"):
    ppr = paragraph._p.get_or_add_pPr()
    el = ppr.find(qn("w:jc"))
    if el is None:
        el = OxmlElement("w:jc")
        ppr.append(el)
    el.set(qn("w:val"), jc)


def add_paragraph(doc, text, size=BODY_SIZE, bold=False, align="thaiDistribute",
                  space_before=0, space_after=2, indent_left=0, line_spacing=1.15,
                  math=True, color=None, italic=False):
    """เพิ่มย่อหน้า รองรับ inline markup + (option) สมการ LaTeX->OMML

    color=None -> ดำ (พฤติกรรมเดิม); ส่ง BLUE/RED เพื่อทำบทเสียง/เฉลย
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if indent_left:
        p.paragraph_format.left_indent = Mm(indent_left)
    set_para_align(p, align)
    if text:
        for kind, val, seg_bold, vert in render_spans(text, math):
            b = bold or seg_bold
            if kind == "math":
                put_math(p, val, size, b, color, italic)
                continue
            run = p.add_run(val)
            set_run_font(run, size, bold=b, color=color, italic=italic)
            _set_vert_align(run, vert)
    return p


def add_image_paragraph(doc, img_path, width_mm, align="center",
                        space_before=2, space_after=2, indent_left=0):
    """ฝังรูปเป็นย่อหน้าเดี่ยว กว้าง width_mm (สูงย่อตามสัดส่วนอัตโนมัติ)

    คืนค่า paragraph. run ที่ได้มี w:drawing และ text ว่าง —
    apply_thai_linebreak() มี guard `if new != run.text` จึงไม่ไปล้าง drawing ทิ้ง
    (ห้ามแก้ guard นั้นเป็นการ assign ตรง ๆ เด็ดขาด)
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if indent_left:
        p.paragraph_format.left_indent = Mm(indent_left)
    set_para_align(p, align)
    run = p.add_run()
    run.add_picture(img_path, width=Mm(width_mm))
    return p


def fill_paragraph(paragraph, content, size=BODY_SIZE, base_bold=False,
                   align="thaiDistribute", math=True):
    """เขียนข้อความลงย่อหน้าที่มีอยู่แล้ว — รองรับ inline markup **และสมการ** เหมือน add_paragraph

    เดิมฟังก์ชันนี้ไม่มี math path ทำให้ `$...$` ในเซลล์ตารางค้างเป็นตัวอักษรดิบ
    และ `	imes` ถูกตีความเป็น tab — แผนการสอนทั้งฉบับอยู่ในตาราง จึงเขียนสมการไม่ได้เลย
    """
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.0
    set_para_align(paragraph, align)
    text = str(content)
    for kind, val, bold, vert in render_spans(text, math):
        b = base_bold or bold
        if kind == "math":
            put_math(paragraph, val, size, b, None, False)
            continue
        run = paragraph.add_run(val)
        set_run_font(run, size, bold=b)
        _set_vert_align(run, vert)


def _set_cell_no_fill(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "auto")


def write_cell(cell, value, base_bold=False, align="thaiDistribute"):
    """เขียนค่าลงเซลล์ — value เป็น str / list[str]"""
    items = value if isinstance(value, list) else [value]
    for extra_p in cell.paragraphs[1:]:
        extra_p._p.getparent().remove(extra_p._p)
    base_p = clear_paragraph(cell.paragraphs[0])
    first = True
    for item in items:
        p = base_p if first else cell.add_paragraph()
        first = False
        fill_paragraph(p, str(item), base_bold=base_bold, align=align)
    _set_cell_no_fill(cell)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "000000")
        borders.append(e)
    tbl_pr.append(borders)


def _set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    tr_pr.append(th)


def add_two_col_table(doc, rows, label_bold=True, header=("หัวข้อ", "รายละเอียด")):
    """ตาราง 2 คอลัมน์ (หน้าปก/แผน): แถวหัวกึ่งกลางตัวหนา ซ้ำทุกหน้า + แถวข้อมูลชิดซ้าย"""
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    hdr = table.add_row()
    write_cell(hdr.cells[0], header[0], base_bold=True, align="center")
    write_cell(hdr.cells[1], header[1], base_bold=True, align="center")
    _set_repeat_header(hdr)
    for label, val in rows:
        tr = table.add_row().cells
        write_cell(tr[0], label, base_bold=label_bold, align="left")
        write_cell(tr[1], val, align="left")
    for r in table.rows:
        r.cells[0].width = Mm(45)
        r.cells[1].width = Mm(140)
    return table


def add_content_table(doc, header, rows, widths_mm=None, col_align=None):
    """ตารางเนื้อหา (หน้า 2+) คอลัมน์อิสระ — AutoFit to window, หัวซ้ำทุกหน้า

    widths_mm  ระบุความกว้างคงที่ต่อคอลัมน์ (มม.) — ใช้เมื่อ autofit ให้ผลไม่ดี
               เช่น คอลัมน์เลขข้อที่มีอักขระเดียวแต่ถูกยืดจนกว้างเกินจำเป็น
    col_align  จัดวางต่อคอลัมน์ เช่น ("center","center","thaiDistribute")
    """
    ncol = len(header)
    table = doc.add_table(rows=0, cols=ncol)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    tbl_pr = table._tbl.tblPr
    tblW = OxmlElement("w:tblW")
    layout = OxmlElement("w:tblLayout")
    if widths_mm:
        table.autofit = False
        tblW.set(qn("w:w"), "0")
        tblW.set(qn("w:type"), "auto")
        layout.set(qn("w:type"), "fixed")
    else:
        tblW.set(qn("w:w"), "5000")      # AutoFit to window 100%
        tblW.set(qn("w:type"), "pct")
        layout.set(qn("w:type"), "autofit")
    tbl_pr.append(tblW)
    tbl_pr.append(layout)
    hdr = table.add_row()
    for i, h in enumerate(header):
        write_cell(hdr.cells[i], h, base_bold=True, align="center")
    _set_repeat_header(hdr)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            a = col_align[i] if col_align and i < len(col_align) else "left"
            write_cell(cells[i], val, align=a)
    if widths_mm:
        # ต้องตั้งที่ทุกเซลล์ ไม่ใช่ที่ column — python-docx เขียน w:tcW รายเซลล์
        for r in table.rows:
            for i, w in enumerate(widths_mm[:ncol]):
                r.cells[i].width = Mm(w)
    return table


def add_heading(doc, text, size=HEADING_SIZE, align="center", space_before=6, space_after=10):
    return add_paragraph(doc, text, size=size, bold=True, align=align,
                         space_before=space_before, space_after=space_after, math=False)


def configure_section(doc, header_text):
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Mm(20)
    sec.bottom_margin = Mm(20)
    sec.left_margin = Mm(25)
    sec.right_margin = Mm(20)

    hp = clear_paragraph(sec.header.paragraphs[0])
    set_para_align(hp, "right")
    r = hp.add_run(header_text)
    set_run_font(r, HF_SIZE)

    fp = clear_paragraph(sec.footer.paragraphs[0])
    set_para_align(fp, "center")
    run = fp.add_run()
    set_run_font(run, HF_SIZE)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def set_normal_style(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE
    style.font.color.rgb = BLACK
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT_NAME)


def _zwsp_text(text):
    if not text or _th_tokenize is None or ZWSP in text or not _THAI_RE.search(text):
        return text
    return ZWSP.join(_th_tokenize(text, engine="newmm", keep_whitespace=True))


def _iter_all_paragraphs(doc):
    def from_tables(tables):
        for t in tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p
                    yield from from_tables(cell.tables)
    yield from doc.paragraphs
    yield from from_tables(doc.tables)
    for sec in doc.sections:
        for hf in (sec.header, sec.footer):
            for p in hf.paragraphs:
                yield p
            yield from from_tables(hf.tables)


def apply_thai_linebreak(doc):
    """แทรก ZWSP ตามขอบคำ + ตั้ง w:lang bidi=th-TH (ทำหลังเอกสารเสร็จ)"""
    for p in _iter_all_paragraphs(doc):
        for run in p.runs:
            new = _zwsp_text(run.text)
            if new != run.text:
                run.text = new
            rpr = run._element.get_or_add_rPr()
            lang = rpr.find(qn("w:lang"))
            if lang is None:
                lang = OxmlElement("w:lang")
                rpr.append(lang)
            lang.set(qn("w:bidi"), "th-TH")


def new_document(header_text):
    """สร้าง Document ตั้ง normal style + section (A4/header/footer) พร้อมใช้"""
    doc = Document()
    set_normal_style(doc)
    configure_section(doc, header_text)
    return doc


def save(doc, out_path):
    apply_thai_linebreak(doc)
    out_dir = os.path.dirname(out_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
    doc.save(out_path)
