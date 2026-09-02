# -*- coding: utf-8 -*-
"""
verify_docx.py — ตัวทดสอบ (tester) ตรวจโครงสร้าง .docx ของ EDU ONE

ใช้งาน:
    python verify_docx.py <kind> <json> <docx>
    kind = content | lesson_plan | exercise

Exit code: 0 = ผ่านทุก check, 1 = พบปัญหา (แสดง list ก่อน exit), 2 = usage error

ตรวจร่วมทุก kind:
  - A4 portrait (210×297 mm)
  - ทุก run ฟอนต์ TH Sarabun New (latin+complex) + ตั้ง szCs
  - Header section ตรงกับ json["header"]
  - Footer มี field PAGE
  - ไม่มี U+FFFD (replacement char / ข้อมูลหาย)

เฉพาะ content / lesson_plan:
  - มี page break (หน้าปก -> เนื้อหา)
  - จำนวนหน้าเนื้อหา (ไม่นับหน้าปก): content = 3-5 หน้า (บังคับ) ·
    lesson_plan = อย่างน้อย 3 หน้า ยาวเกิน 5 ได้ (แค่ WARN)
    [ใช้ Word COM ถ้ามี; ไม่มีก็ WARN ข้าม]

เฉพาะ exercise (สคีมาใหม่ {BASE}_ex.json — {questions:[...]}):
  - heading "แบบฝึกหัด"/"หน้าเฉลย" bold 16pt center + page break คั่น
  - สแกนแบบ state machine (รองรับบรรทัดรูป/เสียงแทรกระหว่างข้อ — ไม่ใช้ stride คงที่)
  - ทุกข้อมีตัวเลือกครบ ก/ข/ค/ง ตามลำดับ
  - ตัวเลือกที่ isTrue = true ต้องเป็น **สีแดง** (C00000)
  - บรรทัด [เสียง] ต้องเป็น **สีน้ำเงิน** (1F4E9C) หรือแดงถ้าเป็นตัวเลือกที่ถูก
  - จำนวนรูปที่ฝัง + บรรทัด [รูปภาพ] = จำนวนช่องรูปใน JSON
  - หน้าเฉลยเป็นตาราง 3 คอลัมน์ (ข้อ / คำตอบที่ถูก / วิธีคิด) ตรงกับ isTrue ใน JSON
  - หน้าแบบฝึกหัดไม่โผล่ easy/medium/hard/difficulty
  - header เทียบเมื่อส่ง --header มาเท่านั้น (สคีมาใหม่ไม่เก็บ header ใน JSON)
"""
import json
import re
import sys

from docx import Document
from docx.oxml.ns import qn

FONT_NAME = "TH Sarabun New"
CHOICE_LETTERS = ("ก", "ข", "ค", "ง")
CHOICE_PREFIXES = tuple(letter + "." for letter in CHOICE_LETTERS)
ANSWER_LETTERS = set(CHOICE_LETTERS)
ZWSP = "​"
REPLACEMENT = "�"
# ต้องตรงกับ docx_common.RED / docx_common.BLUE
RED_HEX = "C00000"
BLUE_HEX = "1F4E9C"


def strip_zwsp(doc):
    paras = list(doc.paragraphs)
    for sec in doc.sections:
        paras.extend(sec.header.paragraphs)
    for p in paras:
        for run in p.runs:
            if ZWSP in run.text:
                run.text = run.text.replace(ZWSP, "")


def has_page_break(paragraph):
    for r in paragraph._p.iter(qn("w:br")):
        if r.get(qn("w:type")) == "page":
            return True
    return False


def run_font_ok(run):
    rpr = run._element.find(qn("w:rPr"))
    if rpr is None:
        return False
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        return False
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        if rfonts.get(qn(attr)) != FONT_NAME:
            return False
    return rpr.find(qn("w:szCs")) is not None


def para_align(paragraph):
    ppr = paragraph._p.find(qn("w:pPr"))
    if ppr is None:
        return None
    jc = ppr.find(qn("w:jc"))
    return jc.get(qn("w:val")) if jc is not None else None


def footer_has_page_field(section):
    for fp in section.footer.paragraphs:
        for instr in fp._p.iter(qn("w:instrText")):
            if instr.text and "PAGE" in instr.text:
                return True
    return False


def _common_checks(doc, expected_header, failures):
    sec = doc.sections[0]
    # A4
    w = round(sec.page_width.mm) if sec.page_width else 0
    h = round(sec.page_height.mm) if sec.page_height else 0
    if not (w == 210 and h == 297):
        failures.append(f"ขนาดหน้าไม่ใช่ A4 (ได้ {w}×{h} mm)")
    # font
    bad = total = 0
    for p in doc.paragraphs:
        for r in p.runs:
            if not r.text:
                continue
            total += 1
            if not run_font_ok(r):
                bad += 1
    if bad:
        failures.append(f"พบ run ฟอนต์ไม่ใช่ {FONT_NAME}/ไม่ตั้ง szCs: {bad}/{total}")
    # header (exercise ส่ง header มาทาง --header; ถ้าไม่ส่งมาก็ข้ามการเทียบ)
    header_text = " ".join(p.text for p in sec.header.paragraphs).replace(ZWSP, "").strip()
    if expected_header is None:
        if not header_text:
            failures.append("Header ว่าง")
        else:
            print(f"WARN: ไม่ได้ส่ง --header มาเทียบ (header ในไฟล์: {header_text!r})",
                  file=sys.stderr)
    elif header_text != expected_header:
        failures.append(f"Header ไม่ตรง\n  expected: {expected_header!r}\n  got     : {header_text!r}")
    # footer PAGE
    if not footer_has_page_field(sec):
        failures.append("Footer ไม่มี field PAGE")
    # replacement char
    alltext = "".join(p.text for p in doc.paragraphs)
    if REPLACEMENT in alltext:
        failures.append("พบ U+FFFD (replacement char) — ข้อมูลเสียหาย")
    return total


def _word_com_content_pages(docx_path):
    """คืนจำนวนหน้าจริง (Word COM) หรือ None ถ้าไม่มี pywin32/Word"""
    try:
        import os
        import win32com.client  # type: ignore
    except Exception:
        return None
    word = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        d = word.Documents.Open(os.path.abspath(docx_path), ReadOnly=True)
        d.Repaginate()
        pages = d.ComputeStatistics(2)  # wdStatisticPages
        d.Close(False)
        return int(pages)
    except Exception:
        return None
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass


def verify_content_like(kind, json_path, docx_path):
    with open(json_path, "r", encoding="utf-8-sig") as f:
        data = json.load(f)
    doc = Document(docx_path)
    failures = []
    _common_checks(doc, data["header"], failures)

    # page break (cover -> content)
    if not any(has_page_break(p) for p in doc.paragraphs):
        failures.append("ไม่พบ page break ระหว่างหน้าปกและเนื้อหา")

    # ---- จำนวนหน้า (Word COM ถ้ามี) ------------------------------------
    # ★ ขอบบนบังคับเฉพาะ `content` — CLAUDE.md ข้อ 4 กำหนด 3–5 หน้าไว้เป็นข้อบังคับ
    #   ส่วน `lesson_plan` ยาวเกิน 5 ได้ เพราะแผนมี 9 แถวที่ต้องมีทั้งจุดประสงค์
    #   กิจกรรมรายขั้น สื่อ และเกณฑ์การวัด — ยิ่งสายโซ่ครบยิ่งยาว การบีบให้อยู่ใน
    #   5 หน้าจึงเป็นการลงโทษแผนที่ทำถูก (เกิดจริงกับ M1-Math_U1_1_1 ที่ต้องเสีย
    #   รอบแก้ไปตัด หลัง checkwork สั่งเพิ่มเกณฑ์การวัดจนโตเป็น 6 หน้า)
    #   ยังเตือนอยู่ เพื่อไม่ให้เสียสัญญาณว่า "เล่มนี้ยาวผิดปกติ"
    # ขอบล่าง 3 หน้าบังคับทั้งคู่ — สั้นกว่านั้นมักแปลว่ามีแถวที่ยังว่าง
    hard_max = 5 if kind == "content" else None
    pages = _word_com_content_pages(docx_path)
    if pages is None:
        limit = "3-5" if hard_max else "อย่างน้อย 3"
        print(f"WARN: ข้าม page-count check (ไม่มี pywin32/Word) — ต้องตรวจ {limit} หน้าด้วยตน",
              file=sys.stderr)
    else:
        content_pages = pages - 1
        if content_pages < 3:
            failures.append(f"เนื้อหาสั้นไป {content_pages} หน้า (ต้องอย่างน้อย 3)")
        elif hard_max and content_pages > hard_max:
            failures.append(f"เนื้อหายาวไป {content_pages} หน้า (ต้อง 3-{hard_max})")
        elif content_pages > 5:
            print(f"WARN: {content_pages} หน้า ยาวกว่าปกติ (ส่วนใหญ่ 3-5) "
                  "— ไม่ถือว่าผิด แต่ลองดูว่ามีถ้อยคำซ้ำที่ตัดได้ไหม", file=sys.stderr)

    return failures


def para_colors(paragraph):
    """เซ็ตของสี (hex) ที่ใช้ในย่อหน้านี้"""
    out = set()
    for r in paragraph.runs:
        rpr = r._element.find(qn("w:rPr"))
        c = rpr.find(qn("w:color")) if rpr is not None else None
        if c is not None:
            out.add((c.get(qn("w:val")) or "").upper())
    return out


def _expected_media_slots(questions):
    """นับช่องสื่อที่ต้องเรนเดอร์อะไรสักอย่าง (มี URL หรือมีคำบรรยาย)"""
    images = audios = 0
    for q in questions:
        if (q.get("imageUrl") or "").strip() or (q.get("imageAlt") or "").strip():
            images += 1
        if (q.get("audioUrl") or "").strip() or (q.get("audioText") or "").strip():
            audios += 1
        for c in q.get("choices") or []:
            ctype = (c.get("contentType") or "text").strip()
            has = (c.get("content") or "").strip() or (c.get("alt") or "").strip()
            if ctype == "image" and has:
                images += 1
            elif ctype == "audio" and has:
                audios += 1
    return images, audios


def _correct_index(q):
    for i, c in enumerate(q.get("choices") or []):
        if c.get("isTrue"):
            return i
    return -1


def verify_exercise(json_path, docx_path, expected_header=None):
    with open(json_path, "r", encoding="utf-8-sig") as f:
        data = json.load(f)
    questions = data.get("questions")
    if not isinstance(questions, list) or not questions:
        return ["ex.json ไม่มีคีย์ 'questions' ที่เป็น list"]
    n = len(questions)

    doc = Document(docx_path)
    strip_zwsp(doc)
    paras = list(doc.paragraphs)
    failures = []
    _common_checks(Document(docx_path), expected_header, failures)

    idx_ex = next((i for i, p in enumerate(paras) if p.text.strip() == "แบบฝึกหัด"), None)
    idx_ans = next((i for i, p in enumerate(paras) if p.text.strip() == "หน้าเฉลย"), None)
    if idx_ex is None:
        failures.append('ไม่พบ heading "แบบฝึกหัด"')
    if idx_ans is None:
        failures.append('ไม่พบ heading "หน้าเฉลย"')

    for label, idx in (("แบบฝึกหัด", idx_ex), ("หน้าเฉลย", idx_ans)):
        if idx is None:
            continue
        p = paras[idx]
        sizes = [r.font.size.pt for r in p.runs if r.font.size]
        bolds = [bool(r.font.bold) for r in p.runs]
        if not bolds or not all(bolds):
            failures.append(f'heading "{label}" ไม่ Bold ทุก run')
        if not sizes or not all(s == 16 for s in sizes):
            failures.append(f'heading "{label}" ขนาดไม่ใช่ 16pt (ได้ {sizes})')
        if para_align(p) != "center":
            failures.append(f'heading "{label}" ไม่กึ่งกลาง')

    if idx_ex is not None and idx_ans is not None:
        if not any(has_page_break(paras[i]) for i in range(idx_ex, idx_ans)):
            failures.append("ไม่พบ page break ระหว่างแบบฝึกหัดและหน้าเฉลย")

        # ---- state machine: รองรับบรรทัดสื่อแทรกได้ (ไม่ใช้ stride คงที่) ----
        ex_paras = [p for p in paras[idx_ex + 1:idx_ans] if p.text.strip()]
        q_re = re.compile(r"^(\d+)\.\s")
        c_re = re.compile(r"^([ก-ง])\.")
        seen_q = 0
        choice_slots = []          # ตัวอักษรตัวเลือกของข้อปัจจุบัน
        correct_para = {}          # (qno, ci) -> paragraph ของตัวเลือกนั้น
        img_lines = aud_lines = 0
        blue_bad = []

        def close_question():
            if seen_q and choice_slots != list(CHOICE_LETTERS):
                failures.append(
                    f"ข้อ {seen_q}: ตัวเลือกเป็น {choice_slots or '[]'} คาดหวัง ['ก','ข','ค','ง']")

        for p in ex_paras:
            text = p.text.strip()
            mq = q_re.match(text)
            if mq:
                close_question()
                num = int(mq.group(1))
                seen_q += 1
                if num != seen_q:
                    failures.append(f"ลำดับคำถามเพี้ยน: เจอ {num} คาดหวัง {seen_q}")
                choice_slots = []
                continue
            mc = c_re.match(text)
            if mc and seen_q:
                choice_slots.append(mc.group(1))
                correct_para[(seen_q, len(choice_slots) - 1)] = p
                continue
            # บรรทัดสื่อ / legend
            if text.startswith("[รูปภาพ]"):
                img_lines += 1
            elif text.startswith("[เสียง]"):
                aud_lines += 1
                cols = para_colors(p)
                if not (BLUE_HEX in cols or RED_HEX in cols):
                    blue_bad.append(text[:40])
        close_question()

        if seen_q != n:
            failures.append(f"นับคำถามได้ {seen_q} ข้อ คาดหวัง {n}")
        for t in blue_bad:
            failures.append(f'บรรทัดเสียงไม่ใช่สีน้ำเงิน/แดง: "{t}"')

        # ---- ตัวเลือกที่ถูกต้องเป็นสีแดง ----
        for qi, q in enumerate(questions, start=1):
            ci = _correct_index(q)
            if ci < 0:
                failures.append(f"ข้อ {qi}: JSON ไม่มีตัวเลือก isTrue = true")
                continue
            if sum(1 for c in q.get("choices") or [] if c.get("isTrue")) > 1:
                failures.append(f"ข้อ {qi}: JSON มี isTrue = true มากกว่า 1 ตัว")
            p = correct_para.get((qi, ci))
            if p is None:
                failures.append(f"ข้อ {qi}: ไม่พบย่อหน้าตัวเลือกที่ถูก (ตำแหน่ง {ci + 1})")
            elif RED_HEX not in para_colors(p):
                failures.append(f"ข้อ {qi}: ตัวเลือกที่ถูกไม่ได้เป็นสีแดง")

        # ---- จำนวนสื่อ: รูปที่ฝัง + บรรทัด placeholder ต้องเท่าช่องที่ต้องเรนเดอร์ ----
        exp_img, exp_aud = _expected_media_slots(questions)
        embedded = len(doc.inline_shapes)
        if embedded + img_lines != exp_img:
            failures.append(
                f"ช่องรูปไม่ครบ: ฝัง {embedded} + placeholder {img_lines} "
                f"= {embedded + img_lines} คาดหวัง {exp_img}")
        if aud_lines != exp_aud:
            failures.append(f"บรรทัดเสียงมี {aud_lines} คาดหวัง {exp_aud}")

        # ---- ห้ามรั่วระดับความยาก (เฉลยสีแดงเป็นความตั้งใจ จึงไม่เช็คคำว่าเฉลย) ----
        ex_text = " ".join(p.text for p in paras[idx_ex + 1:idx_ans]).lower()
        for kw in ("easy", "medium", "hard", "difficulty"):
            if kw in ex_text:
                failures.append(f'หน้าแบบฝึกหัดมีคำว่า "{kw}" โผล่')

    # ---- หน้าเฉลย: ตาราง 3 คอลัมน์ ข้อ / คำตอบที่ถูก / วิธีคิด ----
    # ★ ต้องอ่านจาก doc.tables ไม่ใช่ doc.paragraphs — python-docx ไม่นับย่อหน้า
    #   ที่อยู่ในเซลล์ตารางเข้าไปใน doc.paragraphs ถ้าอ่านผิดที่ ตัวตรวจจะรายงานว่า
    #   "ไม่มีหน้าเฉลย" ทั้งที่มีอยู่ครบ
    if idx_ans is not None:
        # strip_zwsp() แตะเฉพาะ doc.paragraphs — ข้อความในเซลล์ตารางยังมี ZWSP
        # ที่ apply_thai_linebreak แทรกไว้ ต้องลอกออกก่อนเทียบ ไม่งั้นไม่มีวันตรงกัน
        def cell(c):
            return c.text.replace(ZWSP, "").strip()

        want_head = ["ข้อ", "คำตอบที่ถูก", "วิธีคิด"]
        table = None
        for t in doc.tables:
            head = [cell(c) for c in t.rows[0].cells] if t.rows else []
            if head[:3] == want_head:
                table = t
                break
        if table is None:
            failures.append("ไม่พบตารางหน้าเฉลย (หัวตาราง: %s)" % " / ".join(want_head))
        else:
            body = table.rows[1:]
            if len(body) != n:
                failures.append(f"ตารางเฉลยมี {len(body)} แถว คาดหวัง {n}")
            for i, row in enumerate(body, start=1):
                cells = [cell(c) for c in row.cells]
                if not cells[0].isdigit() or int(cells[0]) != i:
                    failures.append(f"ตารางเฉลยแถว {i}: เลขข้อ = {cells[0]!r} คาดหวัง {i}")
                    continue
                q = questions[i - 1]
                ci = _correct_index(q)
                want = CHOICE_LETTERS[ci] if 0 <= ci < len(CHOICE_LETTERS) else "?"
                if cells[1] != want:
                    failures.append(
                        f"ตารางเฉลยข้อ {i} = {cells[1]!r} ไม่ตรง JSON = {want}")
                if (q.get("solutionSteps") or "").strip() and not cells[2].strip("—- "):
                    failures.append(f"ตารางเฉลยข้อ {i}: ช่องวิธีคิดว่าง")

    return failures


def main():
    argv = sys.argv[1:]
    expected_header = None
    if "--header" in argv:
        i = argv.index("--header")
        if i + 1 >= len(argv):
            print("--header ต้องตามด้วยข้อความ", file=sys.stderr)
            sys.exit(2)
        expected_header = argv[i + 1]
        argv = argv[:i] + argv[i + 2:]

    if len(argv) != 3:
        print("usage: python verify_docx.py <content|lesson_plan|exercise> <json> <docx> "
              "[--header <ข้อความ>]", file=sys.stderr)
        sys.exit(2)
    kind = argv[0]
    if kind in ("content", "lesson_plan"):
        failures = verify_content_like(kind, argv[1], argv[2])
    elif kind == "exercise":
        failures = verify_exercise(argv[1], argv[2], expected_header=expected_header)
    else:
        print(f"unknown kind: {kind!r}", file=sys.stderr)
        sys.exit(2)

    if failures:
        print(f"FAIL — ตรวจพบ {len(failures)} ปัญหา:")
        for i, f in enumerate(failures, 1):
            print(f"  {i}. {f}")
        sys.exit(1)
    print(f"OK — ผ่านทุก check ({kind})")
    sys.exit(0)


if __name__ == "__main__":
    main()
